#!/usr/bin/env python3
"""时政分析日报 — 国内政策/人事 + 国际地缘/外交，学术框架分析"""

import json, os, re, subprocess, sys
from datetime import datetime
from collections import defaultdict, Counter

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np

for name in ("Heiti TC", "PingFang SC", "STHeiti", "Arial Unicode MS"):
    try:
        fm.findfont(name, fallback_to_default=False)
        plt.rcParams["font.sans-serif"] = [name, "DejaVu Sans"]
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False

UA = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
REFERER = "https://news.sina.com.cn/"
CHART_DIR = "/tmp/political_charts"
os.makedirs(CHART_DIR, exist_ok=True)

C_RED = "#D32F2F"; C_BLUE = "#1976D2"; C_GRAY = "#616161"; C_ORANGE = "#E65100"; C_GREEN = "#2E7D32"

# ═══════════════════════════════
# DATA LAYER
# ═══════════════════════════════

def curl(url):
    r = subprocess.run(
        ["curl", "-4", "-sk", "--connect-timeout", "10",
         "-H", f"User-Agent: {UA}", "-H", f"Referer: {REFERER}", url],
        capture_output=True, timeout=20
    )
    if r.returncode != 0:
        return None
    raw = r.stdout
    for enc in ("utf-8", "gbk", "gb2312", "gb18030"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")

def fetch_sina_news(lid, num=60):
    """Fetch news from Sina feed by channel lid."""
    all_news = []
    for page in range(1, (num // 20) + 3):
        url = f"https://feed.mix.sina.com.cn/api/roll/get?pageid=153&lid={lid}&k=&num=20&page={page}&r=0.1&callback="
        raw = curl(url)
        if not raw:
            break
        try:
            data = json.loads(raw)
            items = data.get("result", {}).get("data", [])
            if not items:
                break
            all_news.extend(items)
            if len(all_news) >= num:
                break
        except json.JSONDecodeError:
            break
    return all_news[:num]

def fetch_sina_search(keywords, per_kw=10):
    """Search Sina News API for specific political keywords."""
    results = []
    seen_urls = set()
    for kw in keywords:
        url = f"https://search.sina.com.cn/api/news?q={kw}&size={per_kw}"
        raw = curl(url)
        if not raw:
            continue
        try:
            data = json.loads(raw)
            items = data.get("data", {}).get("list", [])
            for item in items:
                url_key = item.get("url", "")
                if url_key in seen_urls:
                    continue
                seen_urls.add(url_key)
                title = item.get("title", "")
                title = re.sub(r'<[^>]+>', '', title)  # strip HTML tags
                results.append({
                    "title": title,
                    "ctime": str(item.get("ctime", "")),
                    "url": url_key,
                    "intro": item.get("intro", ""),
                })
        except (json.JSONDecodeError, KeyError):
            continue
    return results

# ═══════════════════════════════
# ANALYSIS ENGINE
# ═══════════════════════════════

def classify_domestic(news_items):
    """Classify domestic news into: 政策法规, 人事调动, 重要会议, 经济政策, 社会治理, 军事国防.
    First filters out clearly international news."""
    # Comprehensive international filter
    intl_filter = [
        # 美洲
        "美国", "特朗普", "拜登", "白宫", "五角大楼", "华盛顿", "美方", "美媒", "美元",
        "加拿大", "墨西哥", "巴西", "阿根廷", "智利", "委内瑞拉", "古巴", "哥伦比亚",
        "秘鲁", "玻利维亚", "乌拉圭", "巴拉圭", "厄瓜多尔", "洪都拉斯", "巴拿马",
        # 欧洲
        "北约", "欧盟", "欧洲", "欧元", "布鲁塞尔",
        "俄罗斯", "普京", "克里姆林宫", "俄方", "俄媒", "莫斯科",
        "乌克兰", "基辅", "泽连斯基",
        "英国", "伦敦", "英国政府", "英媒", "苏纳克", "斯塔默",
        "法国", "巴黎", "马克龙", "爱丽舍宫", "法媒",
        "德国", "柏林", "朔尔茨", "德媒", "德意志",
        "意大利", "罗马", "梅洛尼", "梵蒂冈",
        "西班牙", "葡萄牙", "荷兰", "比利时", "瑞士", "瑞典", "挪威", "丹麦", "芬兰",
        "波兰", "捷克", "斯洛伐克", "匈牙利", "罗马尼亚", "保加利亚",
        "奥地利", "爱尔兰", "冰岛", "希腊", "克罗地亚", "塞尔维亚",
        "波罗的海", "爱沙尼亚", "拉脱维亚", "立陶宛",
        # 中东
        "以色列", "巴勒斯坦", "哈马斯", "加沙", "约旦河西岸", "以军", "以方",
        "伊朗", "德黑兰", "伊朗核", "伊核",
        "沙特", "阿联酋", "卡塔尔", "科威特", "阿曼", "巴林", "也门", "胡塞",
        "伊拉克", "叙利亚", "黎巴嫩", "约旦", "埃及",
        # 亚太
        "日本", "东京", "日本政府", "日方", "日媒", "岸田", "丰田", "索尼",
        "韩国", "首尔", "韩媒", "韩方", "尹锡悦", "三星", "现代",
        "朝鲜", "平壤", "朝方", "朝中社", "金正恩",
        "印度", "新德里", "莫迪", "印方", "印媒",
        "澳大利亚", "新西兰", "堪培拉",
        "东盟", "ASEAN",
        "菲律宾", "马尼拉", "越南", "河内", "泰国", "曼谷", "印尼", "雅加达",
        "马来西亚", "新加坡", "缅甸", "柬埔寨", "老挝", "文莱", "东帝汶",
        "蒙古", "尼泊尔", "孟加拉", "斯里兰卡", "巴基斯坦", "阿富汗",
        # 中亚/高加索
        "哈萨克", "乌兹别克", "塔吉克", "吉尔吉斯", "土库曼",
        "阿塞拜疆", "亚美尼亚", "格鲁吉亚",
        # 非洲
        "非洲", "埃及", "南非", "尼日利亚", "埃塞俄比亚", "肯尼亚", "苏丹",
        "刚果", "安哥拉", "坦桑尼亚", "加纳", "摩洛哥", "阿尔及利亚",
        "突尼斯", "利比亚", "津巴布韦", "索马里",
        # 国际组织
        "联合国", "联合国安理会", "古特雷斯",
        "G20", "G7", "世行", "IMF", "WTO", "APEC", "OPEC",
        "世界卫生", "世界银行", "国际货币",
        # 外国公司/品牌
        "英伟达", "波音", "空客", "苹果", "谷歌", "微软", "特斯拉",
        "OpenAI", "Meta", "亚马逊",
        # 跨国组织
        "南共市", "非盟", "阿盟", "海合会",
    ]

    categories = {
        "政策法规": [],
        "人事调动": [],
        "重要会议": [],
        "经济政策": [],
        "社会治理": [],
        "军事国防": [],
        "其他": [],
    }
    kw_map = {
        "政策法规": [
            "政策", "法规", "条例", "办法", "规定", "制度", "改革", "方案", "规划",
            "意见", "通知", "印发", "出台", "修订", "施行", "立法", "执法", "行政",
            "许可", "审批", "监管", "整治", "专项", "试点", "示范",
        ],
        "人事调动": [
            "任命", "任免", "免去", "调任", "接任", "当选", "就任", "履新",
            "换届", "提拔", "辞职", "免职", "人事", "干部", "班子", "领导",
            "省委", "市委", "县委", "书记", "省长", "市长", "县长", "部长",
            "局长", "主任", "厅长", "处长", "拟任", "公示", "简历",
        ],
        "重要会议": [
            "会议", "大会", "全会", "座谈会", "峰会", "论坛", "讲话", "致辞",
            "指示", "部署", "强调", "调研", "考察", "主持", "出席", "报告",
            "学习", "传达", "贯彻", "落实", "政协", "人大", "党代会",
        ],
        "经济政策": [
            "经济", "GDP", "财政", "货币", "税收", "产业", "补贴", "基建",
            "投资", "消费", "出口", "债务", "房地产", "金融", "央行", "信贷",
            "利率", "存款", "贷款", "融资", "上市", "股市", "债券", "基金",
            "国企", "民营", "中小", "创业", "就业", "收入", "物价",
            "发改委", "商务部", "工信部", "科技部", "市场监管",
        ],
        "社会治理": [
            "教育", "医疗", "社保", "养老", "住房", "环保", "生态",
            "安全", "生产", "事故", "舆论", "网络", "数据", "隐私",
            "交通", "铁路", "公路", "民航", "水运",
            "乡村", "农村", "城市", "社区", "街道",
            "文化", "旅游", "体育", "文物", "非遗",
            "公安", "检察", "法院", "司法", "律师",
            "高考", "中考", "招生", "学位",
        ],
        "军事国防": [
            "军队", "国防", "军事", "战区", "演习", "装备", "武器",
            "导弹", "军舰", "空军", "海军", "陆军", "火箭军", "战略",
            "服役", "退役", "征兵", "官兵", "军校", "军演", "实战化",
            "南海", "东海", "台海", "钓鱼岛", "藏南", "阿克赛钦",
        ],
    }

    domestic_count = 0
    intl_count = 0
    for item in news_items:
        title = item.get("title", "")
        ctime = item.get("ctime", "")
        url = item.get("url", "")

        # Check if this is international news
        is_intl = any(kw in title for kw in intl_filter)
        if is_intl:
            intl_count += 1
            continue  # Skip international news in domestic classification

        domestic_count += 1
        matched = "其他"
        for cat, kws in kw_map.items():
            if any(kw in title for kw in kws):
                matched = cat
                break
        entry = {"title": title, "ctime": ctime, "url": url}
        if matched == "其他":
            matched = "社会治理"
        categories[matched].append(entry)

    # print(f"  [domestic filter] kept {domestic_count}, filtered {intl_count} intl")
    return categories

def classify_international(news_items):
    """Classify international news into: 大国博弈, 地缘冲突, 多边关系, 经济治理, 区域热点"""
    categories = {
        "大国博弈": [],
        "地缘冲突": [],
        "多边关系": [],
        "经济治理": [],
        "区域热点": [],
        "其他": [],
    }
    kw_map = {
        "大国博弈": ["美国", "中国", "俄罗斯", "北约", "对抗", "遏制", "制裁", "关税",
                    "贸易战", "科技战", "脱钩", "供应链", "芯片", "竞争", "战略"],
        "地缘冲突": ["冲突", "战争", "军事", "导弹", "空袭", "袭击", "占领", "停火",
                    "谈判", "和平", "威胁", "核", "武器", "军队"],
        "多边关系": ["联合国", "G20", "G7", "欧盟", "东盟", "上合", "金砖", "APEC",
                    "峰会", "合作", "协议", "协定", "联合声明", "访问", "外长", "元首"],
        "经济治理": ["经济", "通胀", "加息", "降息", "衰退", "增长", "贸易", "投资",
                    "债务", "IMF", "世行", "WTO", "供应链", "能源", "粮食"],
        "区域热点": ["中东", "欧洲", "亚洲", "非洲", "拉美", "印太", "南海", "台海",
                    "朝鲜", "伊朗", "乌克兰", "巴以", "印度"],
    }
    for item in news_items:
        title = item.get("title", "")
        ctime = item.get("ctime", "")
        url = item.get("url", "")
        matched = "其他"
        for cat, kws in kw_map.items():
            if any(kw in title for kw in kws):
                matched = cat
                break
        entry = {"title": title, "ctime": ctime, "url": url}
        if matched == "其他":
            matched = "区域热点"
        categories[matched].append(entry)
    return categories

def extract_key_policies(domestic_cats):
    """Extract and summarize key policy signals from domestic news."""
    policy_items = domestic_cats.get("政策法规", []) + domestic_cats.get("经济政策", [])
    signals = []
    for item in policy_items[:15]:
        title = item["title"]
        # Extract issuing body
        bodies = ["国务院", "发改委", "财政部", "央行", "证监会", "银保监", "工信部",
                  "科技部", "商务部", "外交部", "国防部", "教育部", "卫健委"]
        issuer = "相关部门"
        for b in bodies:
            if b in title:
                issuer = b
                break
        signals.append({"title": title, "issuer": issuer, "ctime": item["ctime"]})
    return signals

def extract_personnel_changes(domestic_cats):
    """Extract personnel changes from domestic news."""
    return domestic_cats.get("人事调动", [])

def extract_meetings(domestic_cats):
    """Extract important meetings and speeches."""
    return domestic_cats.get("重要会议", [])

def analyze_regional_distribution(international_cats):
    """Analyze which regions are most discussed in international news."""
    regions = {
        "北美": ["美国", "加拿大", "墨西哥", "北美", "华盛顿", "白宫", "五角大楼"],
        "欧洲": ["欧盟", "欧洲", "英国", "法国", "德国", "意大利", "北约", "布鲁塞尔"],
        "中东": ["中东", "伊朗", "以色列", "沙特", "巴勒斯坦", "叙利亚", "伊拉克", "也门", "胡塞"],
        "亚太": ["日本", "韩国", "朝鲜", "印度", "澳大利亚", "东盟", "南海", "台海", "印太"],
        "欧亚": ["俄罗斯", "乌克兰", "白俄", "哈萨克", "中亚"],
        "拉美": ["巴西", "阿根廷", "委内瑞拉", "古巴", "墨西哥", "拉美"],
        "非洲": ["非洲", "埃及", "南非", "尼日利亚", "埃塞"],
    }
    all_titles = []
    for cat in international_cats.values():
        all_titles.extend([item["title"] for item in cat])

    counts = {}
    for region, kws in regions.items():
        count = sum(1 for t in all_titles if any(kw in t for kw in kws))
        counts[region] = count
    return counts

def analyze_global_power_dynamics(international_cats):
    """Analyze major power relationships from news patterns."""
    powers = {
        "美国": ["美国", "特朗普", "白宫", "五角大楼", "华盛顿", "美方", "拜登"],
        "中国": ["中国", "北京", "中方", "中共", "我国"],
        "俄罗斯": ["俄罗斯", "普京", "克里姆林宫", "俄方", "莫斯科"],
        "欧盟": ["欧盟", "欧洲", "德国", "法国", "英国"],
    }
    all_titles = []
    for cat in international_cats.values():
        all_titles.extend([item["title"] for item in cat])

    dynamics = {}
    for name, kws in powers.items():
        count = sum(1 for t in all_titles if any(kw in t for kw in kws))
        dynamics[name] = count
    return dynamics

def generate_domestic_outlook(domestic_cats):
    """Generate domestic political outlook with academic framing."""
    outlooks = []

    # Policy direction
    policy_count = len(domestic_cats.get("政策法规", []))
    econ_count = len(domestic_cats.get("经济政策", []))
    if policy_count > 3:
        outlooks.append({
            "title": "政策密集期判断",
            "view": "政策窗口期",
            "analysis": (
                f"近24小时内有{policy_count}条政策法规类新闻和{econ_count}条经济政策信号，"
                "政策发布密度较高。从政治学议程设定理论（Agenda-Setting Theory）视角来看，"
                "政策窗口期的出现通常与重大会议前后、经济周期拐点或重大事件驱动相关。"
                "当前信号表明政策层正在积极回应经济结构性问题，后续需关注配套细则的出台节奏。"
            ),
            "horizon": "1-2周",
        })

    # Personnel
    personnel_count = len(domestic_cats.get("人事调动", []))
    if personnel_count > 0:
        outlooks.append({
            "title": "人事调整动向",
            "view": "持续关注",
            "analysis": (
                f"监测到{personnel_count}条人事任免相关新闻。"
                "根据组织行为学理论，人事调动的频率和层级可以作为判断治理体系运转状态的重要信号。"
                "跨部门/跨地区的干部交流通常反映中央对特定领域的重视程度提升；"
                "而集中性的反腐/问责类人事变动则需关注对政策执行效率的短期影响。"
            ),
            "horizon": "持续跟踪",
        })

    # Meeting cycle
    meeting_count = len(domestic_cats.get("重要会议", []))
    if meeting_count > 3:
        outlooks.append({
            "title": "会议周期与决策节奏",
            "view": "信号积极",
            "analysis": (
                f"监测到{meeting_count}条重要会议/讲话新闻。"
                "在中国的政策制定体制中，重要会议是政策信号的集中释放渠道。"
                "高层讲话中反复出现的关键词（如'高质量发展''新质生产力''安全'等）"
                "往往预判下一阶段的政策优先级。从政治沟通理论看，"
                "频繁的调研和讲话意味着重大决策前的政治共识正在形成。"
            ),
            "horizon": "1-3个月",
        })

    return outlooks

def generate_international_outlook(international_cats, region_dist, power_dynamics):
    """Generate international political outlook with academic IR theory framing."""
    outlooks = []

    # Great power competition
    us_count = power_dynamics.get("美国", 0)
    cn_count = power_dynamics.get("中国", 0)
    ru_count = power_dynamics.get("俄罗斯", 0)

    if us_count + cn_count > 10:
        outlooks.append({
            "title": "中美战略竞争态势",
            "view": "竞争持续",
            "analysis": (
                f"在{us_count + cn_count}条大国博弈相关新闻中，中美关系仍是核心议题。"
                "从国际关系现实主义理论视角来看，守成大国与崛起大国的结构性矛盾（修昔底德陷阱）"
                "在当前表现为：技术脱钩、供应链重构、金融制裁、意识形态竞争四个维度的同步推进。"
                "短期内双方在气候、公共卫生等低政治领域仍存在有限合作空间，"
                "但在半导体、AI、稀土等战略领域，竞争将持续深化。"
                "从自由主义制度理论看，多边机制（G20/联合国）作为矛盾缓冲器的功能正在被削弱。"
            ),
            "horizon": "6-12个月",
        })

    if ru_count > 5:
        outlooks.append({
            "title": "欧亚地缘安全格局",
            "view": "战略僵持",
            "analysis": (
                f"俄罗斯相关新闻{ru_count}条，反映欧亚安全格局仍是国际关注焦点。"
                "从进攻性现实主义视角分析，大国在欧亚大陆的安全困境短期内难以破解。"
                "核威慑作为终极保障的逻辑被强化，常规力量层面的博弈向灰色地带（网络、太空、"
                "代理人冲突）转移。中国在欧亚安全格局中的角色面临重新定位——"
                "既要维护与俄罗斯的战略协作，又要避免被卷入直接冲突。"
            ),
            "horizon": "3-6个月",
        })

    # Regional hotspots
    top_regions = sorted(region_dist.items(), key=lambda x: x[1], reverse=True)[:3]
    hotspot_text = "、".join(f"{r}({c}条)" for r, c in top_regions if c > 0)
    if hotspot_text:
        outlooks.append({
            "title": "区域热点研判",
            "view": "多点紧张",
            "analysis": (
                f"国际新闻区域分布显示：{hotspot_text}。"
                "从地缘政治学的空间分布理论看，当前国际冲突呈多点并发特征，而非冷战时期的单一"
                "对抗轴心。这种'碎片化的紧张'使得全球治理体系的协调成本大幅上升。"
                "对中国而言，周边安全环境面临'东紧西缓、海强陆弱'的态势——"
                "台海/南海方向战略压力最大，中亚/南亚方向相对可控但需要持续经营。"
            ),
            "horizon": "1-3个月",
        })

    # Global governance
    multi_count = len(international_cats.get("多边关系", []))
    if multi_count > 3:
        outlooks.append({
            "title": "全球治理体系演变",
            "view": "加速重构",
            "analysis": (
                f"多边外交相关新闻{multi_count}条，反映全球治理体系处于深度调整期。"
                "从建构主义理论视角看，国际规范（Norms）正在经历新一轮的竞争和重塑——"
                "西方主导的自由主义国际秩序受到'全球南方'和新兴大国的双重挑战。"
                "金砖扩员、上合扩容、'一带一路'多边化等趋势表明，"
                "中国正在从'规则接受者'向'规则塑造者'转型。"
                "但这一转型面临规范兼容性、制度有效性和联盟凝聚力的三重考验。"
            ),
            "horizon": "6-12个月",
        })

    return outlooks

# ═══════════════════════════════
# CHARTS
# ═══════════════════════════════

def chart_domestic_categories(domestic_cats):
    """Bar chart: domestic news category distribution."""
    fig, ax = plt.subplots(figsize=(8, 4.5))
    cats = ["政策法规", "人事调动", "重要会议", "经济政策", "社会治理", "军事国防"]
    counts = [len(domestic_cats.get(c, [])) for c in cats]
    colors = [C_RED, C_ORANGE, C_BLUE, C_GREEN, C_GRAY, "#455A64"]
    bars = ax.bar(cats, counts, color=colors, edgecolor="white", linewidth=0.5)
    for bar, v in zip(bars, counts):
        if v > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3, str(v),
                    ha="center", fontsize=9, fontweight="bold")
    ax.set_title("国内时政新闻分布", fontsize=14, fontweight="bold", pad=12)
    ax.set_ylabel("新闻数量")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/domestic_cats.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

def chart_international_categories(international_cats):
    """Bar chart: international news category distribution."""
    fig, ax = plt.subplots(figsize=(8, 4.5))
    cats = ["大国博弈", "地缘冲突", "多边关系", "经济治理", "区域热点"]
    counts = [len(international_cats.get(c, [])) for c in cats]
    colors = ["#1565C0", C_RED, C_GREEN, C_ORANGE, "#6A1B9A"]
    bars = ax.bar(cats, counts, color=colors, edgecolor="white", linewidth=0.5)
    for bar, v in zip(bars, counts):
        if v > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3, str(v),
                    ha="center", fontsize=9, fontweight="bold")
    ax.set_title("国际时政新闻分布", fontsize=14, fontweight="bold", pad=12)
    ax.set_ylabel("新闻数量")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/intl_cats.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

def chart_region_distribution(region_dist):
    """Horizontal bar: regional distribution of international coverage."""
    fig, ax = plt.subplots(figsize=(7, 4))
    items = sorted(region_dist.items(), key=lambda x: x[1], reverse=True)
    regions = [i[0] for i in items]
    counts = [i[1] for i in items]
    colors = plt.cm.RdYlGn([c/max(max(counts),1) for c in counts])
    ax.barh(regions, counts, color=colors, edgecolor="white")
    for i, v in enumerate(counts):
        if v > 0:
            ax.text(v + 0.3, i, str(v), va="center", fontsize=9, fontweight="bold")
    ax.set_title("国际新闻区域热度", fontsize=14, fontweight="bold", pad=12)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/region_dist.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

def chart_power_dynamics(power_dynamics):
    """Pie/Donut-like horizontal bar: major power media attention."""
    fig, ax = plt.subplots(figsize=(6, 3.5))
    items = sorted(power_dynamics.items(), key=lambda x: x[1], reverse=True)
    labels = [i[0] for i in items]
    values = [i[1] for i in items]
    colors = ["#1565C0", C_RED, "#455A64", "#2E7D32"]
    bars = ax.barh(labels, values, color=colors[:len(labels)], edgecolor="white")
    for bar, v in zip(bars, values):
        if v > 0:
            ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                    str(v), va="center", fontsize=10, fontweight="bold")
    ax.set_title("大国关系关注度", fontsize=14, fontweight="bold", pad=12)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/power_dynamics.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

def chart_timeline(categories, title, filename):
    """Timeline of news count by date."""
    from collections import Counter
    all_ctimes = []
    for items in categories.values():
        for item in items:
            ctime = item.get("ctime", "")
            if ctime:
                try:
                    ts = int(ctime)
                    d = datetime.fromtimestamp(ts).strftime("%Y-%m-%d")
                    all_ctimes.append(d)
                except (ValueError, OSError):
                    pass
    if not all_ctimes:
        return None
    date_counts = Counter(all_ctimes)
    dates = sorted(date_counts.keys())[-14:]
    counts = [date_counts[d] for d in dates]

    fig, ax = plt.subplots(figsize=(8, 3))
    ax.fill_between(range(len(dates)), counts, alpha=0.3, color=C_BLUE)
    ax.plot(range(len(dates)), counts, marker="o", color=C_BLUE, linewidth=2, markersize=4)
    ax.set_xticks(range(len(dates)))
    ax.set_xticklabels([d[5:] for d in dates], rotation=45, fontsize=7)
    ax.set_title(title, fontsize=13, fontweight="bold", pad=12)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/{filename}.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path

# ═══════════════════════════════
# MAIN
# ═══════════════════════════════

print("=" * 60)
print("时政分析日报生成器 v1.0")
print("=" * 60)

# ── Fetch ──
print("\n[1/4] 获取时政新闻...")
# 2510 = 国内, 2511 = 国际, 2669 = 社会
domestic_raw = fetch_sina_news(lid=2510, num=80)
intl_raw = fetch_sina_news(lid=2511, num=80)
# Supplement: keyword search for policy & personnel
kw_domestic = fetch_sina_search(["国务院", "发改委", "人事任免", "习近平", "央行", "外交"], 10)
print(f"  国内: {len(domestic_raw)} 条 | 国际: {len(intl_raw)} 条 | 关键词补充: {len(kw_domestic)} 条")

if not domestic_raw and not intl_raw:
    print("ERROR: 无法获取新闻数据")
    sys.exit(1)

# ── Analyze ──
print("\n[2/4] 执行分析...")
# Merge feed + keyword search for domestic
domestic_all = list(domestic_raw)
domestic_all.extend(kw_domestic)
# Deduplicate by URL
seen = set()
domestic_dedup = []
for item in domestic_all:
    u = item.get("url", "")
    if u not in seen:
        seen.add(u)
        domestic_dedup.append(item)
domestic_cats = classify_domestic(domestic_dedup)
intl_cats = classify_international(intl_raw)
key_policies = extract_key_policies(domestic_cats)
personnel = extract_personnel_changes(domestic_cats)
meetings = extract_meetings(domestic_cats)
region_dist = analyze_regional_distribution(intl_cats)
power_dynamics = analyze_global_power_dynamics(intl_cats)
domestic_outlook = generate_domestic_outlook(domestic_cats)
intl_outlook = generate_international_outlook(intl_cats, region_dist, power_dynamics)

dom_total = sum(len(v) for v in domestic_cats.values())
intl_total = sum(len(v) for v in intl_cats.values())
print(f"  国内分类: {', '.join(f'{k}({len(v)})' for k, v in domestic_cats.items() if len(v) > 0)}")
print(f"  国际分类: {', '.join(f'{k}({len(v)})' for k, v in intl_cats.items() if len(v) > 0)}")
print(f"  政策信号: {len(key_policies)} 条 | 人事: {len(personnel)} 条 | 会议: {len(meetings)} 条")

# ── Deep Analysis ──
def analyze_policy_pattern(policy_items):
    """Analyze policy news for patterns: issuing bodies, policy areas, regulatory direction."""
    if not policy_items:
        return None
    bodies = Counter()
    areas = Counter()
    for item in policy_items:
        title = item["title"]
        for b in ["国务院", "发改委", "财政部", "央行", "证监会", "工信部", "科技部",
                   "商务部", "外交部", "国防部", "教育部", "卫健委", "生态环境",
                   "交通运输", "农业农村", "文化和旅游", "中央办公厅", "深改委"]:
            if b in title:
                bodies[b] += 1
        for a in ["经济", "金融", "产业", "科技", "环保", "教育", "医疗", "住房",
                   "土地", "税收", "外贸", "基建", "农业", "能源", "数据", "AI"]:
            if a in title:
                areas[a] += 1
    top_bodies = bodies.most_common(5)
    top_areas = areas.most_common(5)
    direction = "扩张性"
    tighten_kw = ["限制", "监管", "整治", "查处", "禁止", "严控", "收紧", "清理"]
    expand_kw = ["支持", "鼓励", "促进", "放宽", "试点", "补贴", "减免", "优化"]
    tight = sum(1 for item in policy_items if any(kw in item["title"] for kw in tighten_kw))
    expand = sum(1 for item in policy_items if any(kw in item["title"] for kw in expand_kw))
    if expand > tight:
        direction = "支持/鼓励型政策为主"
    elif tight > expand:
        direction = "监管/规范型政策为主"
    else:
        direction = "政策信号均衡"
    return {
        "top_bodies": top_bodies,
        "top_areas": top_areas,
        "direction": direction,
        "total": len(policy_items),
        "tight_count": tight,
        "expand_count": expand,
    }

def analyze_personnel_pattern(personnel_items):
    """Analyze personnel changes: level, region, type."""
    if not personnel_items:
        return None
    levels = {"中央": 0, "省部级": 0, "地市级": 0, "县级": 0, "其他": 0}
    types = {"任命": 0, "免职": 0, "调任": 0, "辞职": 0, "审查": 0, "其他": 0}
    provinces = Counter()
    for item in personnel_items:
        title = item["title"]
        for kw, key in [("省委书记", "省部级"), ("省长", "省部级"), ("部长", "中央"),
                         ("市委书记", "地市级"), ("市长", "地市级"), ("县委书记", "县级"),
                         ("县长", "县级"), ("中央", "中央"), ("国家", "中央")]:
            if kw in title:
                levels[key] += 1
                break
        else:
            levels["其他"] += 1
        for kw, key in [("任命", "任命"), ("免职", "免职"), ("调任", "调任"),
                         ("辞职", "辞职"), ("审查", "审查"), ("调查", "审查")]:
            if kw in title:
                types[key] += 1
                break
        else:
            types["其他"] += 1
        for p in ["山东", "广东", "江苏", "浙江", "河南", "四川", "湖北", "湖南",
                   "河北", "福建", "安徽", "辽宁", "陕西", "江西", "广西", "云南",
                   "贵州", "山西", "吉林", "黑龙江", "甘肃", "内蒙古", "新疆",
                   "西藏", "海南", "宁夏", "青海", "北京", "上海", "天津", "重庆"]:
            if p in title:
                provinces[p] += 1
                break
    return {
        "levels": levels,
        "types": types,
        "top_provinces": provinces.most_common(5),
        "total": len(personnel_items),
    }

def analyze_international_dynamics(intl_cats, region_dist, power_dynamics):
    """Deep analysis of international situation with multi-theory framing."""
    analysis = {}

    # Power competition assessment
    us_count = power_dynamics.get("美国", 0)
    cn_count = power_dynamics.get("中国", 0)
    ru_count = power_dynamics.get("俄罗斯", 0)
    eu_count = power_dynamics.get("欧盟", 0)
    if us_count > 0:
        us_share = us_count / max(sum(power_dynamics.values()), 1) * 100
        analysis["power_focus"] = (
            f"美国以{us_count}条相关新闻占据国际关注度的{us_share:.0f}%，"
            f"中国相关{cn_count}条。从国际体系理论看，当前国际议程设置权仍高度集中于美国。"
        )
    top_region = sorted(region_dist.items(), key=lambda x: x[1], reverse=True)[0] if region_dist else None
    if top_region:
        analysis["region_focus"] = (
            f"国际新闻区域分布中，{top_region[0]}以{top_region[1]}条居首，"
            f"反映该地区地缘政治热度最高。区域热点的集中度可通过赫芬达尔指数衡量——"
            f"集中度越高，说明国际冲突越趋于局部化而非全局扩散。"
        )

    # Conflict type analysis
    conflict_items = intl_cats.get("地缘冲突", [])
    if conflict_items:
        c_types = Counter()
        for item in conflict_items:
            title = item["title"]
            for ct, kws in [("军事对抗", ["军事", "导弹", "空袭", "打击", "进攻"]),
                            ("外交危机", ["外交", "抗议", "召见", "驱逐"]),
                            ("安全威胁", ["恐袭", "爆炸", "极端", "武装"]),
                            ("领土争端", ["领土", "边界", "主权", "争议"])]:
                if any(kw in title for kw in kws):
                    c_types[ct] += 1
                    break
            else:
                c_types["其他"] += 1
        analysis["conflict_types"] = dict(c_types.most_common(3))

    return analysis

policy_analysis = analyze_policy_pattern(domestic_cats.get("政策法规", []))
personnel_analysis = analyze_personnel_pattern(personnel)
intl_analysis = analyze_international_dynamics(intl_cats, region_dist, power_dynamics)

# ── Charts ──
print("\n[3/4] 生成图表...")
charts = {
    "domestic_cats": chart_domestic_categories(domestic_cats),
    "intl_cats": chart_international_categories(intl_cats),
    "region_dist": chart_region_distribution(region_dist),
    "power_dynamics": chart_power_dynamics(power_dynamics),
}
tl_dom = chart_timeline(domestic_cats, "国内时政新闻时间线", "timeline_dom")
tl_intl = chart_timeline(intl_cats, "国际时政新闻时间线", "timeline_intl")
if tl_dom:
    charts["tl_dom"] = tl_dom
if tl_intl:
    charts["tl_intl"] = tl_intl
print(f"  {len(charts)} 张图表已生成")

# ═══════════════════════════════
# BUILD DOCX
# ═══════════════════════════════

print("\n[4/4] 生成报告文档...")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def fmt_date(ctime_str):
    try:
        ts = int(ctime_str)
        return datetime.fromtimestamp(ts).strftime("%Y-%m-%d")
    except (ValueError, OSError):
        return ""

def add_news_bullet(doc, item):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item["title"])
    run.font.size = Pt(10)
    d = fmt_date(item.get("ctime", ""))
    if d:
        run2 = p.add_run(f'  ({d})')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(150, 150, 150)

def add_analysis_box(doc, text):
    """Add an indented analysis paragraph with distinct styling."""
    p = doc.add_paragraph()
    run = p.add_run(f'▸ 分析：{text}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(int(C_BLUE[1:3],16), int(C_BLUE[3:5],16), int(C_BLUE[5:7],16))
    p.paragraph_format.left_indent = Cm(0.8)

# ── COVER ──
title = doc.add_heading('时政分析日报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'{datetime.now().strftime("%Y年%m月%d日")} | 基于学术框架的国际国内时政深度分析')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

# ── Executive Summary ──
doc.add_heading('内容摘要', level=1)

# Synthesize a real summary paragraph
summary_lines = []
if policy_analysis:
    summary_lines.append(
        f"本日国内政策信号以{policy_analysis['direction']}，"
        f"主要发声机构为{'、'.join(b for b,_ in policy_analysis['top_bodies'][:3])}，"
        f"重点领域涉及{'、'.join(a for a,_ in policy_analysis['top_areas'][:3])}。"
    )
if personnel_analysis and personnel_analysis["total"] > 0:
    pl = personnel_analysis
    top_types = [f"{k}({v})" for k, v in pl["types"].items() if v > 0][:3]
    summary_lines.append(
        f"人事层面监测到{pl['total']}条变动，类型为{'、'.join(top_types)}，"
        f"主要涉及{'、'.join(p for p,_ in pl['top_provinces'][:3]) if pl['top_provinces'] else '多省份'}。"
    )
if intl_analysis.get("power_focus"):
    summary_lines.append(intl_analysis["power_focus"])
if intl_analysis.get("region_focus"):
    summary_lines.append(intl_analysis["region_focus"])
if intl_analysis.get("conflict_types"):
    ct = intl_analysis["conflict_types"]
    summary_lines.append(f"国际冲突类型分布：{'；'.join(f'{k}({v}条)' for k,v in ct.items())}。")

for line in summary_lines:
    doc.add_paragraph(line, style='List Bullet')

# ═══════════════════════════════
# PART ONE: DOMESTIC
# ═══════════════════════════════

doc.add_page_break()
doc.add_heading('第一部分：国内时政', level=1)

# S1: Policy & Regulations
doc.add_heading('一、政策法规动态', level=2)
policy_items = domestic_cats.get("政策法规", [])
if policy_items:
    doc.add_paragraph(f'近24小时监测到{len(policy_items)}条政策法规类新闻。')
    for item in policy_items[:12]:
        add_news_bullet(doc, item)

    # Analysis
    if policy_analysis:
        doc.add_paragraph()
        doc.add_heading('【深度分析】政策格局研判', level=3)
        bodies_str = '、'.join(f'{b}({c}条)' for b, c in policy_analysis['top_bodies'])
        areas_str = '、'.join(f'{a}({c}条)' for a, c in policy_analysis['top_areas'])
        add_analysis_box(doc,
            f"政策发布主体集中度：{bodies_str}。"
            f"重点领域：{areas_str}。"
            f"政策方向：{policy_analysis['direction']}（收紧类{policy_analysis['tight_count']}条 vs 支持类{policy_analysis['expand_count']}条）。"
        )
        if policy_analysis['top_bodies']:
            top_body = policy_analysis['top_bodies'][0][0]
            add_analysis_box(doc,
                f"从官僚政治理论视角看，{top_body}的密集发文表明该部门在当前的跨部门博弈中处于议程推动者位置。"
                f"其政策产出节奏可以作为判断高层对该领域重视程度的先行指标。"
                f"后续应关注国务院常务会议是否将相关议题列入审议日程——"
                f"若列入，则政策将从'部门推动'升格为'国家意志'，执行力度和资源配给将显著提升。"
            )
        if policy_analysis['expand_count'] > policy_analysis['tight_count']:
            add_analysis_box(doc,
                f"当前政策基调偏积极（{policy_analysis['expand_count']}条支持/鼓励 vs {policy_analysis['tight_count']}条监管/限制），"
                f"从政治经济周期理论看，这可能意味着政策层在主动释放积极信号以稳定市场预期。"
                f"但需警惕'政策繁荣'后的执行瓶颈——出台密度高不等于落地效果好，关键看配套细则和考核机制是否跟进。"
            )
        elif policy_analysis['tight_count'] > policy_analysis['expand_count']:
            add_analysis_box(doc,
                f"当前政策以监管规范型为主（{policy_analysis['tight_count']}条收紧 vs {policy_analysis['expand_count']}条支持），"
                f"可能对应特定行业的整顿周期。从历史经验看，规范期通常持续3-6个月，"
                f"之后进入'规范+发展并重'阶段。建议关注整顿对象和整顿边界的明确程度。"
            )
else:
    doc.add_paragraph("本日未监测到重大政策法规发布。")

# S2: Personnel Changes
doc.add_heading('二、人事调动', level=2)
if personnel:
    doc.add_paragraph(f'监测到{len(personnel)}条人事任免相关新闻：')
    for item in personnel[:15]:
        add_news_bullet(doc, item)

    if personnel_analysis:
        doc.add_paragraph()
        doc.add_heading('【深度分析】人事格局研判', level=3)
        # Level distribution
        levels_info = '、'.join(f'{k}{v}人' for k, v in personnel_analysis['levels'].items() if v > 0)
        add_analysis_box(doc,
            f"本次人事变动层级分布：{levels_info}。"
            f"变动类型以{'、'.join(f'{k}({v}条)' for k,v in personnel_analysis['types'].items() if v > 0)}为主。"
        )
        if personnel_analysis['top_provinces']:
            provs = '、'.join(f'{p}({c}条)' for p, c in personnel_analysis['top_provinces'])
            add_analysis_box(doc,
                f"地域分布：{provs}。从组织行为学来看，"
                f"人事调动的区域集中度是观察中央对地方治理关注度分配的关键窗口。"
                f"若某省份短期内密集出现人事调整，通常对应："
                f"（1）该省即将迎来重大政策试点或重大工程；"
                f"（2）中央对该省此前的工作绩效不满意；"
                f"（3）正常的换届周期。需结合当地经济数据和政策动态交叉验证。"
            )
        if personnel_analysis['levels'].get('省部级', 0) > 3:
            add_analysis_box(doc,
                f"省部级变动{personnel_analysis['levels']['省部级']}人，密度偏高。"
                f"省部级干部是全国治理体系的关键节点，其变动频率和方向（晋升/平调/退休/问责）"
                f"是判断政治周期位置的核心指标。大规模省部级调整通常发生在党代会前后，"
                f"若在非换届期出现，需关注是否与特定事件或政策转向相关。"
            )
else:
    doc.add_paragraph("本日未监测到重大人事调动信息。")

# S3: Key Meetings
doc.add_heading('三、重要会议与讲话', level=2)
if meetings:
    doc.add_paragraph(f'监测到{len(meetings)}条重要会议/讲话新闻：')
    for item in meetings[:10]:
        add_news_bullet(doc, item)

    # Meeting analysis
    keywords_in_meetings = Counter()
    for item in meetings:
        for kw in ["高质量", "安全", "改革", "创新", "民生", "开放", "法治", "生态",
                    "数字经济", "新质生产力", "共同富裕", "现代化", "乡村振兴"]:
            if kw in item["title"]:
                keywords_in_meetings[kw] += 1
    if keywords_in_meetings:
        top_kws = keywords_in_meetings.most_common(5)
        kws_str = '、'.join(f'"{k}"({v}次)' for k, v in top_kws)
        add_analysis_box(doc,
            f"会议讲话中的高频政策关键词：{kws_str}。"
            f"从政治沟通理论来看，高层讲话中的关键词频率是政策优先级排序的信号。"
            f"重复出现的关键词大概率对应下一阶段的资源重点配置方向。"
            f"建议追踪这些关键词在未来1-2个月内是否出现在国务院或部委的具体政策文件中——"
            f"若出现，则说明'讲话定调→政策落地'的传导链已启动。"
        )
else:
    doc.add_paragraph("本日未监测到重大会议新闻。")

# S4: Economic Policy
doc.add_heading('四、经济政策信号', level=2)
econ_items = domestic_cats.get("经济政策", [])
if econ_items:
    doc.add_paragraph(f'监测到{len(econ_items)}条经济政策信号：')
    for item in econ_items[:10]:
        add_news_bullet(doc, item)

    # Economic analysis
    econ_kw = Counter()
    for item in econ_items:
        for kw in ["财政", "货币", "房地产", "基建", "消费", "出口", "就业", "中小企业",
                    "制造业", "数字经济", "绿色", "新能源"]:
            if kw in item["title"]:
                econ_kw[kw] += 1
    if econ_kw:
        top_econ = econ_kw.most_common(4)
        add_analysis_box(doc,
            f"经济政策热点集中：{'、'.join(f'{k}({v})' for k,v in top_econ)}。"
            f"从宏观政策分析框架看，当前财政-货币组合的边际变化方向是判断经济政策立场的核心。"
            f"若财政端基建/专项债信号增多而货币端稳健中性，则政策组合偏向'财政发力+货币配合'的"
            f"经典逆周期调节模式；若两者同步宽松，则经济下行压力可能超出公开数据所反映的程度。"
        )
else:
    doc.add_paragraph("本日无显著经济政策信号。")

# S5: Social Governance
doc.add_heading('五、社会治理动态', level=2)
social_items = domestic_cats.get("社会治理", [])
if social_items:
    for item in social_items[:8]:
        add_news_bullet(doc, item)
    if len(social_items) > 5:
        add_analysis_box(doc,
            f"社会治理类新闻{len(social_items)}条，覆盖面较广。"
            f"从治理理论看，社会治理新闻的数量本身即为社会稳定性的代理指标——"
            f"突发安全事件、公共卫生、自然灾害类新闻的集中出现需结合地域分布"
            f"判断是局部偶发还是系统性风险信号。"
        )

# S6: Military & Defense
doc.add_heading('六、军事国防', level=2)
mil_items = domestic_cats.get("军事国防", [])
if mil_items:
    for item in mil_items[:8]:
        add_news_bullet(doc, item)
else:
    doc.add_paragraph("本日无重大军事国防新闻。")

# Domestic chart
if charts.get("domestic_cats"):
    doc.add_paragraph()
    doc.add_paragraph().add_run().add_picture(charts["domestic_cats"], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# S7: Domestic Outlook
doc.add_paragraph()
doc.add_heading('七、国内形势前瞻', level=2)
for outlook in domestic_outlook:
    p_title = doc.add_paragraph()
    run = p_title.add_run(outlook["title"])
    run.bold = True
    run.font.size = Pt(12)
    p_meta = doc.add_paragraph()
    run_v = p_meta.add_run(f'判断: {outlook["view"]}　|　时间维度: {outlook["horizon"]}')
    run_v.font.size = Pt(10)
    run_v.font.color.rgb = RGBColor(int(C_BLUE[1:3],16), int(C_BLUE[3:5],16), int(C_BLUE[5:7],16))
    doc.add_paragraph(outlook["analysis"])

# ═══════════════════════════════
# PART TWO: INTERNATIONAL
# ═══════════════════════════════

doc.add_page_break()
doc.add_heading('第二部分：国际时政', level=1)

# S8: Great Power Competition
doc.add_heading('八、大国博弈', level=2)
great_power = intl_cats.get("大国博弈", [])
if great_power:
    doc.add_paragraph(f'监测到{len(great_power)}条大国博弈相关新闻：')
    for item in great_power[:12]:
        add_news_bullet(doc, item)

    if intl_analysis.get("power_focus"):
        doc.add_paragraph()
        doc.add_heading('【深度分析】大国战略态势', level=3)
        add_analysis_box(doc, intl_analysis["power_focus"])
        us, cn = power_dynamics.get("美国", 0), power_dynamics.get("中国", 0)
        if us > cn * 1.5:
            add_analysis_box(doc,
                f"美国议程设置优势明显（美{us}条 vs 中{cn}条）。"
                f"从国际政治传播学视角看，美国通过其全球媒体网络的议程设置能力，"
                f"仍在定义'什么问题是重要问题'。这种'软权力'不对称——而非单纯的军事或经济实力差距——"
                f"是中美战略竞争中最被低估的维度。中国近年通过CGTN、新华社海外分社等渠道加大投入，"
                f"但改变全球议程设置权的'时滞效应'可能需要5-10年才能充分显现。"
            )
        if power_dynamics.get("俄罗斯", 0) > 5:
            add_analysis_box(doc,
                f"俄罗斯相关新闻{power_dynamics.get('俄罗斯', 0)}条，在大国博弈叙事中占据显著位置。"
                f"这强化了一个结构性判断：当前国际体系正在从'单极时刻后的多极幻想'"
                f"走向'中美俄三角博弈的新常态'。这一三角关系的特殊之处在于——"
                f"中俄在制衡美国霸权上有共同利益，但在欧亚地缘秩序上存在结构性张力。"
                f"中国需要在中俄'肩并肩'的战略协作与对欧、对美关系的弹性空间之间维持微妙平衡。"
            )

# S9: Geopolitical Conflicts
doc.add_heading('九、地缘冲突与安全', level=2)
conflicts = intl_cats.get("地缘冲突", [])
if conflicts:
    for item in conflicts[:10]:
        add_news_bullet(doc, item)
    if intl_analysis.get("conflict_types"):
        add_analysis_box(doc,
            f"冲突类型分布：{'；'.join(f'{k}({v}条)' for k,v in intl_analysis['conflict_types'].items())}。"
            f"从冲突研究（Conflict Studies）视角看，不同类型冲突对中国的传导路径截然不同："
            f"军事对抗类直接影响中国海外利益保护和军贸风险评估；"
            f"外交危机类考验中国作为联合国安理会常任理事国的立场平衡；"
            f"安全威胁类涉及'一带一路'沿线项目的人员安全评估。"
        )
else:
    doc.add_paragraph("本日无重大地缘冲突新闻。")

# S10: Multilateral Relations
doc.add_heading('十、多边关系与外交', level=2)
multilateral = intl_cats.get("多边关系", [])
if multilateral:
    for item in multilateral[:10]:
        add_news_bullet(doc, item)
    if len(multilateral) > 3:
        add_analysis_box(doc,
            f"多边外交新闻{len(multilateral)}条。从自由制度主义理论看，"
            f"多边机制活跃度是国际合作的晴雨表。当前需要区分两种多边主义："
            f"一种是美国主导的'俱乐部式多边主义'（G7、AUKUS、芯片四方联盟），"
            f"另一种是中国推动的'包容性多边主义'（金砖扩员、上合扩容、全球发展倡议）。"
            f"从建构主义视角看，这两种多边主义的竞争本质是关于'21世纪国际规则由谁书写'的规范之争。"
        )

# S11: Global Economic Governance
doc.add_heading('十一、全球经济治理', level=2)
global_econ = intl_cats.get("经济治理", [])
if global_econ:
    for item in global_econ[:10]:
        add_news_bullet(doc, item)

# S12: Regional Hotspots
doc.add_heading('十二、区域热点', level=2)
regional = intl_cats.get("区域热点", [])
if regional:
    for item in regional[:10]:
        add_news_bullet(doc, item)
    top_region_item = sorted(region_dist.items(), key=lambda x: x[1], reverse=True)
    if top_region_item:
        top_r_names = '、'.join(f'{r}({c}条)' for r, c in top_region_item[:3])
        add_analysis_box(doc,
            f"区域关注度排序：{top_r_names}。各区域对中国的影响路径和干预能力不同："
            f"亚太方向直接关系中国核心安全利益和'一带一路'关键节点；"
            f"中东方向影响能源安全和人民币国际化进程；"
            f"欧洲方向决定中欧经贸关系走向和技术合作空间。"
            f"从战略优先级看，中国的外交-安全资源配置应大致遵循'亚太优先、欧亚并举、全球参与'的梯次结构。"
        )

# International charts
doc.add_paragraph()
if charts.get("intl_cats"):
    doc.add_paragraph().add_run().add_picture(charts["intl_cats"], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
if charts.get("region_dist"):
    doc.add_paragraph().add_run().add_picture(charts["region_dist"], width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
if charts.get("power_dynamics"):
    doc.add_paragraph().add_run().add_picture(charts["power_dynamics"], width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# S13: International Outlook
doc.add_paragraph()
doc.add_heading('十三、国际形势前瞻', level=2)
for outlook in intl_outlook:
    p_title = doc.add_paragraph()
    run = p_title.add_run(outlook["title"])
    run.bold = True
    run.font.size = Pt(12)
    p_meta = doc.add_paragraph()
    run_v = p_meta.add_run(f'判断: {outlook["view"]}　|　时间维度: {outlook["horizon"]}')
    run_v.font.size = Pt(10)
    run_v.font.color.rgb = RGBColor(int(C_RED[1:3],16), int(C_RED[3:5],16), int(C_RED[5:7],16))
    doc.add_paragraph(outlook["analysis"])

# ═══ FOOTER ═══
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— 以上分析基于学术框架，纯属研究性质，不构成任何政治立场声明 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(150, 150, 150)
run.italic = True

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run(
    f'数据: 新浪新闻 | 框架: 现实主义/自由主义/建构主义 IR理论 '
    f'| 生成: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(150, 150, 150)

# ── SAVE ──
output_dir = '/Users/wwwqwq/Desktop/项目/时政报告'
os.makedirs(output_dir, exist_ok=True)
output_path = f'{output_dir}/时政分析日报_{datetime.now().strftime("%Y-%m-%d")}.docx'
doc.save(output_path)

print(f'\n{"="*60}')
print(f'报告已生成: {output_path}')
print(f'国内: {dom_total} 条新闻 | 国际: {intl_total} 条新闻')
print(f'政策信号: {len(key_policies)} | 人事: {len(personnel)} | 会议: {len(meetings)}')
print(f'含 {len(charts)} 张图表')
print(f'{"="*60}')
